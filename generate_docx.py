# /// script
# dependencies = [
#   "python-docx",
#   "google-genai",
# ]
# ///

import os
import sys
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document(student_name, lang, questions, answers, filepath):
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    
    # Title / Header
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(
        "VINSCHOOL MATHEMATICS LEARNING PORTAL\n" if lang == 'en' else "HỆ THỐNG HỌC TẬP TOÁN HỌC VINSCHOOL\n"
    )
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(79, 70, 229) # Indigo

    subtitle_run = title_p.add_run(
        "Personalized Homework Worksheet (Cambridge Stage 6)\n" if lang == 'en' else "Phiếu Bài Tập Cá Nhân Hóa (Mạch Cambridge Stage 6)\n"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(100, 116, 139) # Slate

    # Metadata Table
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.autofit = False
    
    col_widths = [Inches(3.25), Inches(3.25)]
    for row in meta_table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    # Fill metadata
    meta_table.cell(0, 0).paragraphs[0].text = f"Student: {student_name}" if lang == 'en' else f"Học sinh: {student_name}"
    meta_table.cell(0, 0).paragraphs[0].runs[0].bold = True
    meta_table.cell(0, 1).paragraphs[0].text = "Class: 6A1" if lang == 'en' else "Lớp: 6A1"
    meta_table.cell(1, 0).paragraphs[0].text = "Subject: Grade 6 Mathematics" if lang == 'en' else "Môn học: Toán Lớp 6"
    meta_table.cell(1, 1).paragraphs[0].text = "Date Generated: 2026-06-16" if lang == 'en' else "Ngày tạo: 16/06/2026"

    # Add spacing
    doc.add_paragraph()

    # Introduction text
    intro_p = doc.add_paragraph()
    intro_p.add_run(
        "Dear student, based on your practice dashboard reflections, please solve these targeted review questions to solidify your core understanding. Show all your calculation steps clearly." if lang == 'en' else
        "Thầy cô chào em, dựa trên chẩn đoán về các lỗi làm bài gần đây, em hãy giải các câu hỏi ôn tập trọng tâm dưới đây để củng cố kiến thức. Hãy trình bày chi tiết các bước tính toán."
    )
    
    doc.add_paragraph()

    # Questions Table
    q_table = doc.add_table(rows=len(questions) + 1, cols=2)
    q_table.autofit = False
    
    # Set headers
    hdr_row = q_table.rows[0]
    hdr_row.cells[0].width = Inches(3.5)
    hdr_row.cells[1].width = Inches(3.0)
    
    set_cell_background(hdr_row.cells[0], "4F46E5")
    set_cell_background(hdr_row.cells[1], "4F46E5")
    
    hdr_p1 = hdr_row.cells[0].paragraphs[0]
    hdr_p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_hdr1 = hdr_p1.add_run("Target Question / Câu Hỏi Ôn Tập" if lang == 'en' else "Câu Hỏi Ôn Tập Trọng Tâm")
    run_hdr1.bold = True
    run_hdr1.font.color.rgb = RGBColor(255, 255, 255)

    hdr_p2 = hdr_row.cells[1].paragraphs[0]
    hdr_p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_hdr2 = hdr_p2.add_run("Your Working Out & Final Answer" if lang == 'en' else "Nháp Trình Bày & Đáp Án Cuối")
    run_hdr2.bold = True
    run_hdr2.font.color.rgb = RGBColor(255, 255, 255)

    # Fill questions
    for idx, q_text in enumerate(questions):
        row = q_table.rows[idx + 1]
        row.cells[0].width = Inches(3.5)
        row.cells[1].width = Inches(3.0)
        set_cell_margins(row.cells[0], top=120, bottom=120, left=120, right=120)
        set_cell_margins(row.cells[1], top=120, bottom=120, left=120, right=120)

        # Draw light grid lines
        set_cell_background(row.cells[0], "F8FAFC" if idx % 2 == 0 else "FFFFFF")
        
        cell_q_p = row.cells[0].paragraphs[0]
        run_num = cell_q_p.add_run(f"Q{idx + 1}. ")
        run_num.bold = True
        cell_q_p.add_run(q_text)

        # Empty paragraph for student space
        cell_a_p = row.cells[1].paragraphs[0]
        cell_a_p.add_run("\n\n\n\n")

    # Spacing before Answers section
    doc.add_page_break()

    # Answers Section Header
    ans_p = doc.add_paragraph()
    ans_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ans_hdr_run = ans_p.add_run("ANSWER KEY & SOCRATIC SOLUTIONS\n" if lang == 'en' else "HƯỚNG DẪN GIẢI CHI TIẾT\n")
    ans_hdr_run.bold = True
    ans_hdr_run.font.size = Pt(13)
    ans_hdr_run.font.color.rgb = RGBColor(16, 185, 129) # Emerald Green

    # List solutions
    for idx, ans_text in enumerate(answers):
        sol_p = doc.add_paragraph()
        run_sol_num = sol_p.add_run(f"Question {idx + 1}: ")
        run_sol_num.bold = True
        sol_p.add_run(ans_text)

    # Footer notice
    doc.add_paragraph()
    foot_p = doc.add_paragraph()
    foot_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    foot_run = foot_p.add_run("Vinschool Cambridge Math Explorer - Dynamic AI Engine")
    foot_run.font.size = Pt(8)
    foot_run.italic = True
    foot_run.font.color.rgb = RGBColor(148, 163, 184)

    doc.save(filepath)
    print(f"Word document saved successfully to: {filepath}")

def get_personalized_content(student_id, lang):
    # Default high-quality fallbacks mapped to standard Stage 6 Gaps
    if lang == 'en':
        questions = [
            "What is the value of the digit 4 in the decimal number 15.648?",
            "Calculate: 3.456 multiplied by 100.",
            "Round the decimal value 8.945 to the nearest tenth.",
            "Write the equivalent fraction of 12/18 in its simplest terms.",
            "If 3 boxes of cookies cost $15, how much do 9 boxes cost?"
        ]
        answers = [
            "4 hundredths (or 4/100, 0.04) because it lies in the second position after the decimal point.",
            "345.6. Multiplying by 100 shifts each digit two columns to the left (decimal point moves 2 places right).",
            "8.9. The digit in the hundredths place is 4, which is less than 5, so we round down.",
            "2/3. Dividing both numerator and denominator by their Greatest Common Divisor (GCD) which is 6.",
            "$45. The scale factor is 9 / 3 = 3. Multiply the original cost by the scale factor: $15 x 3 = $45."
        ]
    else:
        questions = [
            "Chữ số 4 trong số thập phân 15,648 có giá trị là bao nhiêu?",
            "Tính nhẩm kết quả: 3,456 nhân với 100.",
            "Làm tròn số thập phân 8,945 đến hàng phần mười (chữ số thập phân thứ nhất).",
            "Rút gọn phân số 12/18 về dạng phân số tối giản.",
            "Nếu 3 hộp bánh có giá $15, thì 9 hộp bánh như vậy có giá bao nhiêu?"
        ]
        answers = [
            "4 phần trăm (hay 4/100, 0,04) vì chữ số này nằm ở vị trí thứ hai sau dấu phẩy thập phân.",
            "345,6. Khi nhân với 100, mỗi chữ số dịch chuyển sang trái 2 hàng (dấu phẩy dịch sang phải 2 chữ số).",
            "8,9. Chữ số hàng phần trăm là 4 (nhỏ hơn 5) nên ta thực hiện làm tròn xuống.",
            "2/3. Thực hiện chia cả tử số và mẫu số cho ước chung lớn nhất (UCLN) là 6.",
            "$45. Hệ số tỉ lệ số hộp bánh là 9 / 3 = 3. Nhân giá tiền ban đầu với 3: $15 x 3 = $45."
        ]

    # If API key is available, try to prompt Gemini to make them dynamic!
    api_key = os.environ.get("GEMINI_API_KEY")
    if api_key and api_key.strip() != "" and "MY_GEMINI_API_KEY" not in api_key:
        try:
            from google import genai
            client = genai.Client(api_key=api_key)
            
            prompt = f"""
            You are a Cambridge Stage 6 Math Curriculum writer for Vinschool.
            Write 5 personalized math practice questions and detailed step-by-step Socratic answers for student '{student_id}' in language '{lang}'.
            Focus on these areas: decimal place value shift, fraction simplification, and ratio scaling.
            Keep them appropriate for 11-year-olds. Return ONLY a valid JSON object matching this schema:
            {{
                "questions": [
                    "Question 1 text",
                    "Question 2 text",
                    ...
                ],
                "answers": [
                    "Answer 1 explanation",
                    "Answer 2 explanation",
                    ...
                ]
            }}
            Do not wrap the response in markdown blocks. Output raw JSON string only.
            """
            
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt
            )
            
            data = json.loads(response.text.strip())
            if "questions" in data and "answers" in data and len(data["questions"]) == 5:
                return data["questions"], data["answers"]
        except Exception as e:
            print(f"Gemini API dynamic generation bypassed: {e}. Using official Stage 6 curriculum fallbacks.")
            
    return questions, answers

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python generate_docx.py <student_id> <language> <output_path>")
        sys.exit(1)

    student_id = sys.argv[1]
    lang = sys.argv[2]
    output_path = sys.argv[3]
    
    # Retrieve student name mapping
    name_map = {
        'student_minh': 'Hoàng Minh',
        'student_nam': 'Lê Nam',
        'student_mai': 'Nguyễn Mai',
        'student_chloe': 'Chloe Smith',
        'student_duc': 'Trần Đức',
        'student_linh': 'Phạm Linh'
    }
    student_name = name_map.get(student_id, "Vinschool Student")
    
    print(f"Generating personalized DOCX for {student_name}...")
    questions, answers = get_personalized_content(student_id, lang)
    
    # Create directory if not exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    create_document(student_name, lang, questions, answers, output_path)
